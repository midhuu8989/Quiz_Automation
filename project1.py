import os
import streamlit as st
import openai
import json
from io import BytesIO
from docx import Document
from docx.shared import Pt
import pandas as pd
from langchain_openai import ChatOpenAI

# Load OpenAI API key
openai_api_key = st.secrets["OPENAI_API_KEY"]

# Initialize Langchain LLM
llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0, openai_api_key=openai_api_key)

# Predefined subtopics per topic
DEFAULT_SUBTOPICS = {
    "Data Science": ["Data Visualization", "Statistics", "Data Cleaning", "Machine Learning Basics"],
    "Python": ["Loops", "Functions", "Data Types", "OOP"]
}

# Generate MCQs from LLM
def generate_mcqs(subtopic, num_questions):
    prompt = f"""
Generate {num_questions} multiple choice questions on the topic: {subtopic}.
Each question must have 4 options (A–D), and specify the correct answer.
Return as JSON in this format:
[
  {{
    "question": "...",
    "options": {{"A": "...", "B": "...", "C": "...", "D": "..."}},
    "correct": "A"
  }}
]
"""
    result = llm.invoke(prompt)
    return json.loads(result.content)

# Handle duplicates and regenerate them
def remove_and_regenerate_duplicates(all_mcqs):
    seen_questions = set()
    duplicates_info = []
    MAX_RETRIES = 5

    for (subtopic_level, mcq_list) in all_mcqs.items():
        for idx, q in enumerate(mcq_list):
            q_text = q['question'].strip().lower()
            if q_text in seen_questions:
                duplicates_info.append((subtopic_level, idx))
            else:
                seen_questions.add(q_text)

    total_duplicates = len(duplicates_info)

    for (subtopic_level, idx) in duplicates_info:
        sub, level = subtopic_level
        attempts = 0
        while attempts < MAX_RETRIES:
            try:
                new_mcqs = generate_mcqs(sub, 1)
                if not new_mcqs:
                    attempts += 1
                    continue

                new_q = new_mcqs[0]
                new_q_text = new_q['question'].strip().lower()

                if new_q_text not in seen_questions:
                    all_mcqs[subtopic_level][idx] = new_q
                    seen_questions.add(new_q_text)
                    break
            except Exception as e:
                print(f"Retry {attempts+1} failed for {subtopic_level}: {e}")
            attempts += 1

    return all_mcqs, total_duplicates

# Word Export
def create_mcq_doc(sample_mcqs, topic):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    doc.add_heading(f"QUIZ_{topic}_Day9", 0)
    q_counter = 1

    for (subtopic, level), mcqs in sample_mcqs.items():
        for q in mcqs:
            doc.add_paragraph(f"Q{q_counter}_{subtopic}_{level}")
            table = doc.add_table(rows=19, cols=4)
            table.style = 'Table Grid'

            table.cell(0, 0).merge(table.cell(0, 2)).text = f"{q_counter}. {q['question']}({level.title()})"
            table.cell(0, 3).text = "MC"

            table.cell(1, 0).merge(table.cell(1, 2)).text = "Default mark:"
            table.cell(1, 3).text = "1"
            table.cell(2, 0).merge(table.cell(2, 2)).text = "Shuffle the choices?"
            table.cell(2, 3).text = "Yes"
            table.cell(3, 0).merge(table.cell(3, 2)).text = "Number the choices?"
            table.cell(3, 3).text = "A"
            table.cell(4, 0).merge(table.cell(4, 2)).text = "Penalty for each incorrect try:"
            table.cell(4, 3).text = "0"

            table.cell(5, 0).text = "#"
            table.cell(5, 1).text = "Answers"
            table.cell(5, 2).text = "Feedback"
            table.cell(5, 3).text = "Grade"

            for i, opt in enumerate(["A", "B", "C", "D"]):
                r = 6 + i
                table.cell(r, 0).text = opt
                table.cell(r, 1).text = q["options"][opt]
                table.cell(r, 2).text = ""
                table.cell(r, 3).text = "100" if q["correct"] == opt else "0"

            feedback_data = [
                ("", "For any correct response:", "Your answer is correct.", ""),
                ("", "For any incorrect response:", "Your answer is incorrect.", ""),
                ("", "Hint 1:", "", ""),
                ("", "Show the number of correct responses (Hint 1):", "No", ""),
                ("", "Clear incorrect responses (Hint 1):", "No", ""),
                ("", "Tags:", "", ""),
                ("", "", "", ""),
            ]
            for i, (c0, c1, c2, c3) in enumerate(feedback_data):
                r = 11 + i
                table.cell(r, 0).text = c0
                table.cell(r, 1).text = c1
                table.cell(r, 2).text = c2
                table.cell(r, 3).text = c3

            table.cell(18, 0).merge(table.cell(18, 2)).text = "Allows the selection of a single or multiple responses from a pre-defined list. (MC/MA)"
            table.cell(18, 3).text = ""

            doc.add_page_break()
            q_counter += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Excel Export
def save_to_excel(topic, all_mcqs):
    rows = []
    for (subtopic, level), mcqs in all_mcqs.items():
        for q in mcqs:
            row = {
                "Action": "ADD",
                "Question Bank": topic,
                "Sub Bank": subtopic,
                "Difficulty Level": level,
                "Question Instruction": "",
                "Question Text": q["question"],
                "Question Type": "MCQ",
                "Choice1": q["options"]["A"],
                "Choice2": q["options"]["B"],
                "Choice3": q["options"]["C"],
                "Choice4": q["options"]["D"],
                "Choice5": "",
                "Grade1": 1 if q["correct"] == "A" else 0,
                "Grade2": 1 if q["correct"] == "B" else 0,
                "Grade3": 1 if q["correct"] == "C" else 0,
                "Grade4": 1 if q["correct"] == "D" else 0,
                "Grade5": 0,
                "Answer Description": q["options"][q["correct"]]
            }
            rows.append(row)
    df = pd.DataFrame(rows)
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name="Quiz")
    buf.seek(0)
    return buf

# GIFT Format Export
def create_gift_file(sample_mcqs, topic):
    gift_lines = [f"// QUIZ: {topic}\n"]
    q_counter = 1
    for (subtopic, level), mcqs in sample_mcqs.items():
        for q in mcqs:
            gift_lines.append(f"// Q{q_counter}_{subtopic}_{level}")
            question_text = q["question"].replace("\n", " ")
            gift_lines.append(f"::{q_counter}_{subtopic}_{level}:: {question_text} ({level.title()}) {{")
            for opt_key in ["A", "B", "C", "D"]:
                symbol = "=" if opt_key == q["correct"] else "~"
                option_text = q["options"][opt_key].replace("\n", " ")
                gift_lines.append(f"{symbol}{option_text}")
            gift_lines.append("}\n")
            q_counter += 1
    gift_content = "\n".join(gift_lines)
    return BytesIO(gift_content.encode("utf-8"))

# ─────────────── Streamlit UI ───────────────
st.set_page_config(page_title="🧠 MCQ Generator", layout="centered")
st.title("🧠 AI-Powered MCQ Generator")

topic = st.text_input("Enter main topic (e.g., Data Science)")
format_choice = st.radio("Select output format:", ["Word", "Excel", "Moodle GIFT (.txt)"])

custom_subtopics = {}

if topic:
    default_subs = DEFAULT_SUBTOPICS.get(topic, [])
    selected_subs = st.multiselect("Select subtopics:", default_subs)

    st.markdown("---")
    st.subheader("➕ Add Subtopics and Question Difficulty Levels")

    for sub in selected_subs:
        st.markdown(f"### {sub}")
        for level in ["simple", "medium", "complex"]:
            count = st.number_input(f"{level.title()} questions for '{sub}'", min_value=0, max_value=20, value=0, key=f"{sub}_{level}")
            if count > 0:
                custom_subtopics[(sub, level)] = count

    additional_count = st.number_input("Number of additional subtopics to add:", min_value=0, max_value=10, value=0)

    for i in range(additional_count):
        sub_name = st.text_input(f"Custom Subtopic {i+1}", key=f"name_custom_{i}")
        if sub_name:
            for level in ["simple", "medium", "complex"]:
                count = st.number_input(f"{level.title()} questions for '{sub_name}'", min_value=0, max_value=20, value=0, key=f"{sub_name}_{level}_{i}")
                if count > 0:
                    custom_subtopics[(sub_name, level)] = count

    if st.button("🚀 Generate MCQs"):
        all_mcqs = {}
        with st.spinner("Generating questions..."):
            for (sub, level), num in custom_subtopics.items():
                try:
                    all_mcqs[(sub, level)] = generate_mcqs(sub, num)
                except Exception as e:
                    st.error(f"Failed to generate for '{sub}' ({level}): {e}")

        original_mcqs = all_mcqs.copy()
        all_mcqs, dup_count = remove_and_regenerate_duplicates(all_mcqs)

        if dup_count > 0:
            st.warning(f"⚠️ {dup_count} duplicate questions were found and regenerated.")
        else:
            st.success("✅ No duplicate questions found.")

        st.success("✅ MCQs generated!")

        st.markdown("### 📂 Download Original MCQs (May contain duplicates)")
        if format_choice == "Word":
            docx_original = create_mcq_doc(original_mcqs, topic)
            st.download_button("📄 Download Word (With Duplicates)", docx_original, file_name=f"Quiz_{topic}_Original.docx")
        elif format_choice == "Excel":
            xlsx_original = save_to_excel(topic, original_mcqs)
            st.download_button("📊 Download Excel (With Duplicates)", xlsx_original, file_name=f"Quiz_{topic}_Original.xlsx")
        else:
            gift_original = create_gift_file(original_mcqs, topic)
            st.download_button("📑 Download GIFT (With Duplicates)", gift_original, file_name=f"Quiz_{topic}_Original.txt")

        st.markdown("### 🎯 Download Cleaned MCQs (Duplicates Removed & Regenerated)")
        if format_choice == "Word":
            docx_clean = create_mcq_doc(all_mcqs, topic)
            st.download_button("✅ Download Clean Word", docx_clean, file_name=f"Quiz_{topic}_Clean.docx")
        elif format_choice == "Excel":
            xlsx_clean = save_to_excel(topic, all_mcqs)
            st.download_button("✅ Download Clean Excel", xlsx_clean, file_name=f"Quiz_{topic}_Clean.xlsx")
        else:
            gift_clean = create_gift_file(all_mcqs, topic)
            st.download_button("✅ Download Clean GIFT", gift_clean, file_name=f"Quiz_{topic}_Clean.txt")
